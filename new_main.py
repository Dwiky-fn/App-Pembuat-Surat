import pandas as pd
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
from openpyxl import load_workbook
from babel.dates import format_date
from openpyxl.utils.dataframe import dataframe_to_rows

# Atur locale ke bahasa Indonesia

# Fungsi untuk mengatur font dan ukuran teks
def set_font(run, font_name='Times New Roman', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Fungsi untuk mengatur teks menjadi bold
def set_bold(run):
    run.bold = True

def srt_pemberitahuan():
    st.title('Surat Pemberitahuan Kegiatan')

    # Load template surat
    doc = Document('pemberitahuan.docx')

    # Membuat Nomor Surat
    # Input untuk bagian-bagian nomor surat
    with st.expander("Nomor Surat", expanded=False):
        nomor = st.text_input('Nomor Urut', '001')
        panitia = st.text_input('Kepanitiaan (Opsional)', placeholder='Cth: PAN-RAIS')
        bulan_romawi_list = {
            "Januari": "I",
            "Februari": "II",
            "Maret": "III",
            "April": "IV",
            "Mei": "V",
            "Juni": "VI",
            "Juli": "VII",
            "Agustus": "VIII",
            "September": "IX",
            "Oktober": "X",
            "November": "XI",
            "Desember": "XII"
        }
        nama_bulan = st.selectbox('Bulan dalam Angka Romawi', list(bulan_romawi_list.keys()))
        bulan_romawi = bulan_romawi_list[nama_bulan]
        tahun = st.text_input('Tahun', str(datetime.now().year))

        # Format nomor surat
        if panitia:
            no_srt = st.text_input("Nomor Surat", value=f"{nomor}/{panitia}/HMJTE/{bulan_romawi}/{tahun}")
        else:
            no_srt = st.text_input("Nomor Surat", value=f"{nomor}/HMJTE/{bulan_romawi}/{tahun}")
            
    st.write(f"Nomor Surat: {no_srt}")
    
    # Tanggal Pembuatan Surat
    tanggal_input = format_date(datetime.now().date(), format='d MMMM yyyy', locale='id_ID')
    tanggal = st.text_input('Tanggal Pembuatan Surat', value=tanggal_input)

    tujuan_srt = st.radio('Tujuan:', ('KEMENDAGRI BEM', 'Lainnya'))
    if tujuan_srt == 'KEMENDAGRI BEM': tujuan = 'KEMENDAGRI BEM'
    else: tujuan = st.text_input('Tujuan', placeholder='Cth: KEMENDAGRI BEM')

    # Nama Kegiatan
    kegiatan = st.text_input('Kegiatan', placeholder='Cth: Rapat Kerja Kepengurusan HMJ Teknik Elektro')

    # Hari Kegiatan
    opsi_hari = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
    pilih_hari = st.multiselect('Hari:', opsi_hari, placeholder="Pilih Hari Kegiatan")
    hari = ' – '.join(pilih_hari)
    
    # Tanggal Kegiatan
    with st.expander("Tanggal Kegiatan", expanded=False):
        st.write("Pilih tanggal mulai dan selesai:")

        # Pilih tanggal mulai
        tanggal_mulaii = st.date_input('Pilih Tanggal Mulai', value=datetime(2024, 1, 1))
        tanggal_mulai = format_date(tanggal_mulaii, format='d MMMM yyyy', locale='id_ID')
        # Opsi untuk memilih tanggal selesai
        opsi_selesai = st.radio("Atur tanggal selesai?", ("1 Hari", "Lebih dari 1 Hari"))

        if opsi_selesai == "Lebih dari 1 Hari":
            tanggal_selesaii = st.date_input('Pilih Tanggal Selesai', value=datetime(2024, 1, 3))
            tanggal_selesai = format_date(tanggal_selesaii, format='d MMMM yyyy', locale='id_ID') 
            tanggal_keg = st.text_input("",value=f"{tanggal_mulai} – {tanggal_selesai}")
        elif opsi_selesai == "1 Hari":
            tanggal_keg = st.text_input("",value=f"{tanggal_mulai}")

    st.write("Tanggal Kegiatan: ",tanggal_keg)

    # Jam Kegiatan
    with st.expander("Waktu Kegiatan", expanded=False):
        st.write("Waktu Kegiatan")

        # Pilih jam dan menit untuk waktu mulai
        jam_mulai = st.slider('Pilih Jam Mulai', 0, 23, 8)
        menit_mulai = st.slider('Pilih Menit Mulai', 0, 59, 0)

        # Pilih jam dan menit untuk waktu selesai
        opsi_selesai = st.radio("Atur waktu selesai?", ("Ya", "Tidak"))

        if opsi_selesai == "Ya":
            jam_selesai = st.slider('Pilih Jam Selesai', 0, 23, 9)
            menit_selesai = st.slider('Pilih Menit Selesai', 0, 59, 0)
            jam_keg = st.text_input("",value=f"{jam_mulai:02d}.{menit_mulai:02d} – {jam_selesai:02d}.{menit_selesai:02d} WIB")
        else:
            jam_keg = st.text_input("",value=f"{jam_mulai:02d}.{menit_mulai:02d} WIB – Selesai")

    st.write('Waktu Kegiatan:', jam_keg)

    # Tempat Kegiatan
    tempat = st.text_input('Tempat', placeholder='Cth: Gedung Teori TI 4')

    # Memilih Sekretaris
    nama_sekre = st.selectbox('Nama Sekretaris:', ('Desi Viviani', 'Gita Dwi Astuti', 'Dwiky Juniardi', 'Lainnya'))
    if nama_sekre == 'Desi Viviani':
        nim_sekre = '3202216008'
    elif nama_sekre == 'Gita Dwi Astuti':
        nim_sekre = '3202316023'
    elif nama_sekre == 'Dwiky Juniardi':
        nim_sekre = '3202316001'
    else:
        nama_sekre = st.text_input('Nama Sekretaris', placeholder='Nama yang membuat surat')
        nim_sekre = st.text_input('NIM Sekretaris', placeholder='NIM yang membuat surat')
    st.write("Nama: ", nama_sekre)
    st.write("NIM: ",nim_sekre)

    # Tombol submit untuk membuat surat
    submit_button = st.button('Buat Surat')

    if submit_button:
        if no_srt and tanggal and tujuan and kegiatan and tanggal_keg and jam_keg and tempat and nama_sekre and nim_sekre:
            for para in doc.paragraphs:
                if '{no_surat}' in para.text:
                    para.text = para.text.replace('{no_surat}', no_srt)
                    for run in para.runs:
                        set_font(run)
                if '{tanggal}' in para.text:
                    para.text = para.text.replace('{tanggal}', tanggal)
                    for run in para.runs:
                        set_font(run)
                if '{tujuan}' in para.text:
                    for run in para.runs:
                        run.text = run.text.replace('{tujuan}', tujuan)
                        if tujuan in run.text:
                            set_bold(run)
                        set_font(run)
                if '{kegiatan}' in para.text:
                    for run in para.runs:
                        run.text = run.text.replace('{kegiatan}', kegiatan)
                        if tujuan in run.text:
                            set_bold(run)
                        set_font(run)
                if '{hari}' in para.text:
                    para.text = para.text.replace('{hari}', hari)
                    for run in para.runs:
                        set_font(run)
                if '{tanggal_keg}' in para.text:
                    para.text = para.text.replace('{tanggal_keg}', tanggal_keg)
                    for run in para.runs:
                        set_font(run)
                if '{jam_keg}' in para.text:
                    para.text = para.text.replace('{jam_keg}', jam_keg)
                    for run in para.runs:
                        set_font(run)
                if '{tempat}' in para.text:
                    para.text = para.text.replace('{tempat}', tempat)
                    for run in para.runs:
                        set_font(run)

            for table in doc.tables:
                for row in table.rows:
                    cell = row.cells[1]
                    for paragraph in cell.paragraphs:
                        if '{nama_sekre}' in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace('{nama_sekre}', nama_sekre)
                                if nama_sekre in run.text:
                                    set_bold(run)
                                set_font(run)
                        if '{nim_sekre}' in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace('{nim_sekre}', nim_sekre)
                                set_font(run)

        # Membaca file Excel yang sudah ada

        # Simpan dokumen ke BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Tombol untuk mengunduh surat
        st.download_button(
            label = "Simpan Surat",
            data = output.getvalue(),
            file_name = f'{nomor} Pemberitahuan {tujuan}.docx',
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        path_file = "arsip.xlsx"

        data = [{
            "no surat": no_srt,
            "tanggal": tanggal,
            "kegiatan": kegiatan,
            "tujuan": tujuan,
            "perihal": "Pemberitahuan"
        }]

        df_baru = pd.DataFrame(data)
        book = load_workbook(path_file)
        sheet = book.active

        startrow = sheet.max_row
        for r_idx, row in enumerate(dataframe_to_rows(df_baru, index=False, header=False), start=startrow+1):
            for c_idx, value in enumerate(row, start=1):
                sheet.cell(row=r_idx, column=c_idx, value=value)
        book.save(path_file)
        st.success('Data tersimpan')
def arsip_srt():
    st.title("Arsip Surat")
    path_file = "arsip.xlsx"
    df = pd.read_excel(path_file)
    header = {
        "no surat": "Nomor Surat",
        "tanggal": "Tanggal",
        "kegiatan": "Kegiatan",
        "tujuan": "Tujuan",
        "perihal": "Perihal"
    }
    df.rename(columns=header, inplace=True)
    df.index = df.index + 1

    st.table(df)

def main():
    st.sidebar.title('Aplikasi Pembuat Surat Otomatis')

    surat = st.sidebar.radio('Menu', ('Pemberitahuan', 'Peminjaman', 'Arsip'))
    
    if surat == 'Pemberitahuan':
        srt_pemberitahuan()
    elif surat == 'Peminjaman':
        pass  # Fungsi srt_peminjaman bisa ditambahkan
    elif surat == 'Arsip':
        arsip_srt()

if __name__ == '__main__':
    main()
