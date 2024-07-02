import gspread
import streamlit as st
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.oxml.ns import qn
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from oauth2client.service_account import ServiceAccountCredentials

# Fungsi untuk mengatur font dan ukuran teks
def set_font(run, font_name='Times New Roman', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Fungsi untuk mengatur teks menjadi bold
def set_bold(run):
    run.bold = True

# Fungsi untuk membuat surat berdasarkan templat
def srt_pemberitahuan():
    st.title('Surat Pemberitahuan Kegiatan')

    # Load template surat
    doc = Document('pemberitahuan.docx')

    with st.form(key='data'):
        # Input dari pengguna
        no_srt = st.text_input('Nomor Surat', placeholder='001')
        bulan = st.text_input('Bulan', placeholder='Bulan dengan Angka Romawi Cth: VI')
        tanggal = st.text_input('Tanggal', placeholder='Tanggal Pembuatan Surat Cth: 23 Juni 2004')
        
        tujuan_srt = st.radio('Tujuan:', ('KEMENDAGRI BEM', 'Lainnya'))
        st.form_submit_button('Gunakan')
        if tujuan_srt == 'KEMENDAGRI BEM':
            tujuan = 'KEMENDAGRI BEM'
            st.write(tujuan)
        elif tujuan_srt == 'Lainnya':
            tujuan = st.text_input('Tujuan', placeholder='Cth: KEMENDAGRI BEM')
        acara = st.text_input('Kegiatan', placeholder='Cth: Rapat Kerja Kepengurusan HMJ Teknik Elektro')
        hari = st.text_input('Hari Kegiatan', placeholder='Cth: Senin atau Senin - Rabu')
        tanggal_keg = st.text_input('Tanggal Kegiatan', placeholder='Cth: 1 Juli 2024, 1 - 3 Juli 2024, atau 30 Juni 2024 - 2 Juli 2024')
        jam_mulai = st.text_input('Jam Mulai', placeholder='Cth: 08.00')
        jam_selesai = st.text_input('Jam Selesai', placeholder='Cth: 12.00 WIB atau Selesai')
        tempat = st.text_input('Tempat', placeholder='Gedung Teori TI 4')
        nama_sekre = st.radio('Nama Sekretaris:',('Desi Viviani', 'Gita Dwi Astuti', 'Dwiky Juniardi', 'Lainnya'))
        st.form_submit_button('Gunakan Nama')
        if nama_sekre == 'Desi Viviani':
            nim_sekre = '3202216008'
            st.write('NIM Sekretaris:')
            st.write(nim_sekre)
        elif nama_sekre == 'Gita Dwi Astuti':
            nim_sekre = '3202316023'
            st.write('NIM Sekretaris:')
            st.write(nim_sekre)
        elif nama_sekre == 'Dwiky Juniardi':
            nim_sekre = '3202316001'
            st.write('NIM Sekretaris:')
            st.write(nim_sekre)
        elif nama_sekre == 'Lainnya':
            nama_sekre = st.text_input('Nama Sekretaris', placeholder='Nama yang membuat surat')
            nim_sekre = st.text_input('NIM Sekretaris', placeholder='NIM yang membuat surat')
            
        
        submit_button = st.form_submit_button(label='Buat Surat')
    # Ganti placeholder dengan data yang dimasukkan dan atur font
    if submit_button:
        if no_srt and tanggal and tujuan and acara and tanggal_keg and jam_mulai and jam_selesai and tempat and nama_sekre and nim_sekre:
        # Create letter and get result as BytesIO

            for para in doc.paragraphs:
                if '{no_surat}' in para.text:
                    para.text = para.text.replace('{no_surat}', no_srt)
                    for run in para.runs:
                        set_font(run)
                if '{bulan}' in para.text:
                    para.text = para.text.replace('{bulan}', bulan)
                    for run in para.runs:
                        set_font(run)
                if '{tanggal}' in para.text:
                    para.text = para.text.replace('{tanggal}', tanggal)
                    for run in para.runs:
                        set_font(run)
                if '{tujuan}' in para.text:
                    for run in para.runs:
                        run.text = run.text.replace('{tujuan}', tujuan)
                        if tujuan in run.text:
                            set_bold(run)
                        set_font(run)
                if '{kegiatan}' in para.text:
                    for run in para.runs:
                        run.text = run.text.replace('{kegiatan}', acara)
                        if tujuan in run.text:
                            set_bold(run)
                        set_font(run)
                if '{hari}' in para.text:
                    para.text = para.text.replace('{hari}', hari)
                    for run in para.runs:
                        set_font(run)
                if '{tanggal_keg}' in para.text:
                    para.text = para.text.replace('{tanggal_keg}', tanggal_keg)
                    for run in para.runs:
                        set_font(run)
                if '{jam_mulai}' in para.text:
                    para.text = para.text.replace('{jam_mulai}', jam_mulai)
                    for run in para.runs:
                        set_font(run)
                if '{jam_selesai}' in para.text:
                    para.text = para.text.replace('{jam_selesai}', jam_selesai)
                    for run in para.runs:
                        set_font(run)
                if '{tempat}' in para.text:
                    para.text = para.text.replace('{tempat}', tempat)
                    for run in para.runs:
                        set_font(run)

            for table in doc.tables:
                for row in table.rows:
                    cell = row.cells[1]
                    for paragraph in cell.paragraphs:
                        if '{nama_sekre}' in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace('{nama_sekre}', nama_sekre)
                                if nama_sekre in run.text:
                                    set_bold(run)
                                set_font(run)
                        if '{nim_sekre}' in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace('{nim_sekre}', nim_sekre)
                                set_font(run)
        
        # Simpan dokumen ke BytesIO
        output = BytesIO()
        doc.save(output)

        file_name = f'{no_srt} Pemberitahuan {tujuan}.docx'
        data = [no_srt, 'Peminjaman', tujuan, acara, tanggal, file_name]
        worksheet.append_row(data)

        output.seek(0)

        if output:
        # Provide option to download the letter
            st.download_button(
            label = "Unduh Surat",
            data = output.getvalue(),
            file_name = f'{no_srt} Pemberitahuan {tujuan}.docx',
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        upload_word_document(file_name, file_name, "1PgtCbzYcJUjrDfdSZ-L-2GLOo4L1xq3x")
        st.success('Surat Telah di Upload ke Database')

def srt_peminjaman():
    jenis = st.sidebar.radio('Jenis Peminjaman:', ('Ruangan', 'Peralatan'))

    if jenis == 'Peralatan':
        st.title('Surat Peminjaman Peralatan')
        doc = Document('peminjaman_barang.docx')
        
        with st.form(key='data'):
            no_srt = st.text_input('Nomor Surat', placeholder='Cht: 001')
            panitia_keg = st.text_input('Panitia Kegiatan', placeholder='Cth: /PAN-RAIS')
            bulan = st.text_input('Bulan', placeholder='Bulan dengan Angka Romawi Cth: VI')
            tanggal = st.text_input('Tanggal', placeholder='Tanggal Pembuatan Surat Cth: 23 Juni 2004')
            lampiran = st.text_input('Lampiran', placeholder='Cth: 1 atau Berkas')

            tujuan = st.text_input('Tujuan', placeholder='Cth: Kepala Jurusan Teknik Elektro')
            tempat = st.text_input('Tempat Tujuan', placeholder='Politeknik Negeri Pontianak')

            # isi
            kegiatan = st.text_input('Kegiatan', placeholder='Cth: Rapat Kerja Kepengurusan HMJ Teknik Elektro')
            hari = st.text_input('Hari Kegiatan', placeholder='Cth: Senin atau Senin - Rabu')
            tanggal_keg = st.text_input('Tanggal Kegiatan', placeholder='Cth: 1 Juli 2024, 1 - 3 Juli 2024, atau 30 Juni 2024 - 2 Juli 2024')
            jam_mulai = st.text_input('Jam Mulai', placeholder='Cth: 08.00')
            jam_selesai = st.text_input('Jam Selesai', placeholder='Cth: 12.00 WIB atau Selesai')
            tempat_keg = st.text_input('Tempat', placeholder='Ruangan Teori TI 4')

            jab_penanggungjwb = st.text_input('Jabatan Penanggung Jawab', placeholder='Cth: Ketua Panitia')
            nama_penanggungjwb = st.text_input('Nama Penanggung Jawab', placeholder='Cth: Muhammad Ikhwan')
            nim_penanggungjwb = st.text_input('NIM Penanggung Jawab', placeholder='Cth: 3202315001')

            tujuan = st.radio('Tujuan', (
                'Ketua Jurusan Teknik ELektro',
                'Sekretaris Jurusan Teknik Elektro',
                'Koordinator Program Studi Teknik Informatika',
                'Koordinator Program Studi Teknik Listrik',
                'Koordinator Program Studi Teknologi Rekayasa Sistem Elektronika',
                'Kepala Laboratorium Program Studi Teknik Informatika',
                'Kepala Bengkel Program Studi Teknik Listrik',
                'Kepala Laboratorium Program Studi Teknik Listrik',
                'Kepala Laboratorium Teknik Elektronika',
                'Kepala Bengkel Teknik Elektronika',
                'Lainnya'
            ))

            st.form_submit_button(label='Gunakan')

            if tujuan == 'Ketua Jurusan Teknik ELektro':
                nama_tujuan = "Hasan"
                nip_tujuan = '197108201999031003'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Sekretaris Jurusan Teknik Elektro':
                nama_tujuan = "Wiwit Indah Rahayu"
                nip_tujuan = '198802292019032015'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Koordinator Program Studi Teknik Informatika':
                nama_tujuan = "Mariana Syamsudin"
                nip_tujuan = '197503142006042001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Koordinator Program Studi Teknik Listrik':
                nama_tujuan = "Irman"
                nip_tujuan = '196409061990031001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Koordinator Program Studi Teknologi Rekayasa Sistem Elektronika':
                nama_tujuan = "Agus Riyanto"
                nip_tujuan = '197202282006041001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Laboratorium Program Studi Teknik Informatika':
                nama_tujuan = 'Tomi Suryanto'
                nip_tujuan = '199010202019031013'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Bengkel Program Studi Teknik Listrik':
                nama_tujuan = 'Suparno'
                nip_tujuan = '196409131990031002'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Laboratorium Program Studi Teknik Listrik':
                nama_tujuan = 'Wawan Heryawan'
                nip_tujuan = '197010161997021002'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Laboratorium Teknik Elektronika':
                nama_tujuan = 'Medi Yuwono Tharam'
                nip_tujuan = '197007181999031001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Bengkel Teknik Elektronika':
                nama_tujuan = 'M. Ridhwan Sufandi'
                nip_tujuan = '198602042019031013'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Lainnya':
                tujuan = st.text_input('Tujuan Peminjaman', placeholder='Cth: Ketua ')
                nama_tujuan = st.text_input('Nama Tujuan Peminjaman', placeholder='Cth: Hasan')
                nip_tujuan = st.text_input('NIP Tujuan Peminjaman', placeholder='Cth: 197108201999031003')
            # peralatan
            jns_peralatan = jenis
            nama_peralatan = st.text_input('Nama Barang', placeholder='Cth: Proyektor')
            jmlh = st.text_input('Jumlah Barang', placeholder='Cth: 2 Buah')

            submit_button = st.form_submit_button(label='Buat Surat')

        if submit_button:
            if no_srt and panitia_keg and tanggal and tujuan and kegiatan and tanggal_keg and jam_mulai and jam_selesai and tempat_keg and jab_penanggungjwb and nama_penanggungjwb and nim_penanggungjwb and tempat and nama_tujuan and nip_tujuan and jns_peralatan and nama_peralatan and jmlh:
            # Create letter and get result as BytesIO

                for para in doc.paragraphs:
                    if '{no_srt}' in para.text:
                        para.text = para.text.replace('{no_srt}', no_srt)
                        for run in para.runs:
                            set_font(run)
                    if '{panitia_keg}' in para.text:
                        para.text = para.text.replace('{panitia_keg}', panitia_keg)
                        for run in para.runs:
                            set_font(run)
                    if '{bln}' in para.text:
                        para.text = para.text.replace('{bln}', bulan)
                        for run in para.runs:
                            set_font(run)
                    if '{tanggal_srt}' in para.text:
                        para.text = para.text.replace('{tanggal_srt}', tanggal)
                        for run in para.runs:
                            set_font(run)
                    if '{lampiran}' in para.text:
                        para.text = para.text.replace('{lampiran}', lampiran)
                        for run in para.runs:
                            set_font(run)
                    if '{tujuan}' in para.text:
                        for run in para.runs:
                            run.text = run.text.replace('{tujuan}', tujuan)
                            if tujuan in run.text:
                                set_bold(run)
                            set_font(run)
                    if '{tempat}' in para.text:
                        for run in para.runs:
                            run.text = run.text.replace('{tempat}', tempat)
                            if tujuan in run.text:
                                set_bold(run)
                            set_font(run)
                    if '{kegiatan}' in para.text:
                        para.text = para.text.replace('{kegiatan}', kegiatan)
                        for run in para.runs:
                            if tujuan in run.text:
                                set_bold(run)
                            set_font(run)
                    if '{hari_keg}' in para.text:
                        para.text = para.text.replace('{hari_keg}', hari)
                        for run in para.runs:
                            set_font(run)
                    if '{tanggal_keg}' in para.text:
                        para.text = para.text.replace('{tanggal_keg}', tanggal_keg)
                        for run in para.runs:
                            set_font(run)
                    if '{jam_mulai}' in para.text:
                        para.text = para.text.replace('{jam_mulai}', jam_mulai)
                        for run in para.runs:
                            set_font(run)
                    if '{jam_selesai}' in para.text:
                        para.text = para.text.replace('{jam_selesai}', jam_selesai)
                        for run in para.runs:
                            set_font(run)
                    if '{tempat_keg}' in para.text:
                        para.text = para.text.replace('{tempat_keg}', tempat_keg)
                        for run in para.runs:
                            set_font(run)
                    if '{jns_peralatan}' in para.text:
                        para.text = para.text.replace('{jns_peralatan}', jns_peralatan.upper())
                        for run in para.runs:
                            set_font(run)

                for table in doc.tables:
                    for row in table.rows:
                        cell = row.cells[1]
                        for paragraph in cell.paragraphs:
                            if '{jab_penanggungjwb}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{jab_penanggungjwb}', jab_penanggungjwb)
                                    set_font(run)
                            if '{nama_penanggungjwb}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nama_penanggungjwb}', nama_penanggungjwb)
                                    if nama_penanggungjwb in run.text:
                                        set_bold(run)
                                    set_font(run)
                            if '{nim_penanggungjwb}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nim_penanggungjwb}', nim_penanggungjwb)
                                    set_font(run)
                            if '{jmlh}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{jmlh}', jmlh)
                                    set_font(run)
                            if '{tujuan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{tujuan}', tujuan)
                                    set_font(run)
                            if '{nama_tujuan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nama_tujuan}', nama_tujuan)
                                    if nama_tujuan in run.text:
                                        set_bold(run)
                                    set_font(run)
                            if '{nip_tujuan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nip_tujuan}', nip_tujuan)
                                    set_font(run)

                for table in doc.tables:
                    for row in table.rows:
                        cell = row.cells[0]
                        for paragraph in cell.paragraphs:
                            if '{nama_peralatan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nama_peralatan}', nama_peralatan)
                                    set_font(run)

            # Simpan dokumen ke BytesIO
            output = BytesIO()
            doc.save(output)

            file_name = f'{no_srt} Peminjaman {nama_peralatan} {tujuan}.docx'
            data = [no_srt, 'Peminjaman', tujuan, kegiatan, tanggal, file_name]
            worksheet.append_row(data)

            output.seek(0)

            if output:
            # Provide option to download the letter
                st.download_button(
                label = "Unduh Surat",
                data = output.getvalue(),
                file_name = f'{no_srt} Peminjaman {nama_peralatan} {tujuan}.docx',
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            upload_word_document(file_name, file_name, "1PgtCbzYcJUjrDfdSZ-L-2GLOo4L1xq3x")
            st.success('Surat Telah di Upload ke Database')

    elif jenis == 'Ruangan':
        st.title('Surat Peminjaman Ruangan')
        doc = Document('peminjaman_ruangan.docx')

        # no_srt = st.form()
        with st.form(key='data'):
            no_srt = st.text_input('Nomor Surat', placeholder='Cht: 001')
            panitia_keg = st.text_input('Panitia Kegiatan', placeholder='Cth: /PAN-RAIS')
            bulan = st.text_input('Bulan', placeholder='Bulan dengan Angka Romawi Cth: VI')
            tanggal = st.text_input('Tanggal', placeholder='Tanggal Pembuatan Surat Cth: 23 Juni 2004')
            lampiran = st.text_input('Lampiran', placeholder='Cth: 1 atau Berkas')

            tujuan = st.text_input('Tujuan', placeholder='Cth: Kepala Jurusan Teknik Elektro')
            tempat = st.text_input('Tempat Tujuan', placeholder='Politeknik Negeri Pontianak')

            # isi
            kegiatan = st.text_input('Kegiatan', placeholder='Cth: Rapat Kerja Kepengurusan HMJ Teknik Elektro')
            hari = st.text_input('Hari Kegiatan', placeholder='Cth: Senin atau Senin - Rabu')
            tanggal_keg = st.text_input('Tanggal Kegiatan', placeholder='Cth: 1 Juli 2024, 1 - 3 Juli 2024, atau 30 Juni 2024 - 2 Juli 2024')
            jam_mulai = st.text_input('Jam Mulai', placeholder='Cth: 08.00')
            jam_selesai = st.text_input('Jam Selesai', placeholder='Cth: 12.00 WIB atau Selesai')

            jab_penanggungjwb = st.text_input('Jabatan Penanggung Jawab', placeholder='Cth: Ketua Panitia')
            nama_penanggungjwb = st.text_input('Nama Penanggung Jawab', placeholder='Cth: Muhammad Ikhwan')
            nim_penanggungjwb = st.text_input('NIM Penanggung Jawab', placeholder='Cth: 3202315001')

            tujuan = st.radio('Tujuan', (
                'Ketua Jurusan Teknik ELektro',
                'Sekretaris Jurusan Teknik Elektro',
                'Koordinator Program Studi Teknik Informatika',
                'Koordinator Program Studi Teknik Listrik',
                'Koordinator Program Studi Teknologi Rekayasa Sistem Elektronika',
                'Kepala Laboratorium Program Studi Teknik Informatika',
                'Kepala Bengkel Program Studi Teknik Listrik',
                'Kepala Laboratorium Program Studi Teknik Listrik',
                'Kepala Laboratorium Teknik Elektronika',
                'Kepala Bengkel Teknik Elektronika',
                'Lainnya'
            ))

            st.form_submit_button(label='Gunakan')

            if tujuan == 'Ketua Jurusan Teknik ELektro':
                nama_tujuan = "Hasan"
                nip_tujuan = '197108201999031003'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Sekretaris Jurusan Teknik Elektro':
                nama_tujuan = "Wiwit Indah Rahayu"
                nip_tujuan = '198802292019032015'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Koordinator Program Studi Teknik Informatika':
                nama_tujuan = "Mariana Syamsudin"
                nip_tujuan = '197503142006042001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Koordinator Program Studi Teknik Listrik':
                nama_tujuan = "Irman"
                nip_tujuan = '196409061990031001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Koordinator Program Studi Teknologi Rekayasa Sistem Elektronika':
                nama_tujuan = "Agus Riyanto"
                nip_tujuan = '197202282006041001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Laboratorium Program Studi Teknik Informatika':
                nama_tujuan = 'Tomi Suryanto'
                nip_tujuan = '199010202019031013'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Bengkel Program Studi Teknik Listrik':
                nama_tujuan = 'Suparno'
                nip_tujuan = '196409131990031002'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Laboratorium Program Studi Teknik Listrik':
                nama_tujuan = 'Wawan Heryawan'
                nip_tujuan = '197010161997021002'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Laboratorium Teknik Elektronika':
                nama_tujuan = 'Medi Yuwono Tharam'
                nip_tujuan = '197007181999031001'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Kepala Bengkel Teknik Elektronika':
                nama_tujuan = 'M. Ridhwan Sufandi'
                nip_tujuan = '198602042019031013'
                st.write('Nama Tujuan:')
                st.write(nama_tujuan)
                st.write('NIP Tujuan:')
                st.write(nip_tujuan)
            elif tujuan == 'Lainnya':
                tujuan = st.text_input('Tujuan Peminjaman', placeholder='Cth: Ketua ')
                nama_tujuan = st.text_input('Nama Tujuan Peminjaman', placeholder='Cth: Hasan')
                nip_tujuan = st.text_input('NIP Tujuan Peminjaman', placeholder='Cth: 197108201999031003')

            # peralatan
            jns_peralatan = jenis
            nama_peralatan = st.text_input('Nama Ruangan ', placeholder='Cth: Ruangan Teori TI 1')
            jmlh = st.text_input('Jumlah Barang', placeholder='Cth: 1 Kelas')

            submit_button = st.form_submit_button(label='Buat Surat')

        if submit_button:
            if no_srt and panitia_keg and tanggal and lampiran and tujuan and kegiatan and tanggal_keg and jam_mulai and jam_selesai and jab_penanggungjwb and nama_penanggungjwb and nim_penanggungjwb and tempat and nama_tujuan and nip_tujuan and jns_peralatan and nama_peralatan and jmlh:
            # Create letter and get result as BytesIO

                for para in doc.paragraphs:
                    if '{no_srt}' in para.text:
                        para.text = para.text.replace('{no_srt}', no_srt)
                        for run in para.runs:
                            set_font(run)
                    if '{panitia_keg}' in para.text:
                        para.text = para.text.replace('{panitia_keg}', panitia_keg)
                        for run in para.runs:
                            set_font(run)
                    if '{bln}' in para.text:
                        para.text = para.text.replace('{bln}', bulan)
                        for run in para.runs:
                            set_font(run)
                    if '{tanggal_srt}' in para.text:
                        para.text = para.text.replace('{tanggal_srt}', tanggal)
                        for run in para.runs:
                            set_font(run)
                    if '{lampiran}' in para.text:
                        para.text = para.text.replace('{lampiran}', lampiran)
                        for run in para.runs:
                            set_font(run)
                    if '{tujuan}' in para.text:
                        for run in para.runs:
                            run.text = run.text.replace('{tujuan}', tujuan)
                            if tujuan in run.text:
                                set_bold(run)
                            set_font(run)
                    if '{tempat}' in para.text:
                        for run in para.runs:
                            run.text = run.text.replace('{tempat}', tempat)
                            if tujuan in run.text:
                                set_bold(run)
                            set_font(run)
                    if '{kegiatan}' in para.text:
                        para.text = para.text.replace('{kegiatan}', kegiatan)
                        for run in para.runs:
                            if tujuan in run.text:
                                set_bold(run)
                            set_font(run)
                    if '{hari_keg}' in para.text:
                        para.text = para.text.replace('{hari_keg}', hari)
                        for run in para.runs:
                            set_font(run)
                    if '{tanggal_keg}' in para.text:
                        para.text = para.text.replace('{tanggal_keg}', tanggal_keg)
                        for run in para.runs:
                            set_font(run)
                    if '{jam_mulai}' in para.text:
                        para.text = para.text.replace('{jam_mulai}', jam_mulai)
                        for run in para.runs:
                            set_font(run)
                    if '{jam_selesai}' in para.text:
                        para.text = para.text.replace('{jam_selesai}', jam_selesai)
                        for run in para.runs:
                            set_font(run)
                    if '{jns_peralatan}' in para.text:
                        para.text = para.text.replace('{jns_peralatan}', jns_peralatan.upper())
                        for run in para.runs:
                            set_font(run)

                for table in doc.tables:
                    for row in table.rows:
                        cell = row.cells[1]
                        for paragraph in cell.paragraphs:
                            if '{jab_penanggungjwb}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{jab_penanggungjwb}', jab_penanggungjwb)
                                    set_font(run)
                            if '{nama_penanggungjwb}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nama_penanggungjwb}', nama_penanggungjwb)
                                    if nama_penanggungjwb in run.text:
                                        set_bold(run)
                                    set_font(run)
                            if '{nim_penanggungjwb}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nim_penanggungjwb}', nim_penanggungjwb)
                                    set_font(run)
                            if '{jmlh}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{jmlh}', jmlh)
                                    set_font(run)
                            if '{tujuan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{tujuan}', tujuan)
                                    set_font(run)
                            if '{nama_tujuan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nama_tujuan}', nama_tujuan)
                                    if nama_tujuan in run.text:
                                        set_bold(run)
                                    set_font(run)
                            if '{nip_tujuan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nip_tujuan}', nip_tujuan)
                                    set_font(run)

                for table in doc.tables:
                    for row in table.rows:
                        cell = row.cells[0]
                        for paragraph in cell.paragraphs:
                            if '{nama_peralatan}' in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace('{nama_peralatan}', nama_peralatan)
                                    set_font(run)
        
            # Simpan dokumen ke BytesIO
            output = BytesIO()
            doc.save(output)

            file_name = f'{no_srt} Peminjaman {nama_peralatan} {tujuan}.docx'
            data = [no_srt, 'Peminjaman', tujuan, kegiatan, tanggal, file_name]
            # worksheet.append_row(data)

            output.seek(0)

            if output:
            # Provide option to download the letter
                st.download_button(
                label = "Unduh Surat",
                data = output.getvalue(),
                file_name = file_name,
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            upload_word_document(file_name, file_name, "1PgtCbzYcJUjrDfdSZ-L-2GLOo4L1xq3x")
            st.success('Surat Telah di Upload ke Database')

def access_drive():
    '''Mengakses Drive menggunakan API dari Google Cloud '''
    # File path to the service account JSON key file
    SERVICE_ACCOUNT_FILE = 'config/credentials.json'

    # Define the required scopes
    SCOPES = ['https://www.googleapis.com/auth/drive.file']

    # Authenticate and create the service
    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=credentials)

    return service

service = access_drive()

def upload_word_document(file_path, file_name, folder_id=None):
    file_metadata = {'name': file_name}
    if folder_id:
        file_metadata['parents'] = [folder_id]

    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    try:
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        
        print(f'File ID: {file.get("id")}')
        
    except Exception as e:
        print(f'Error uploading file: {e}')

def access_spreadsheet():
    '''Mengakses SpreadSheet menggunakan API dari Google Cloud Console'''
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_name("config/credentials.json", scope)
    client = gspread.authorize(creds)

    try:
        # Mengakses spreadsheet dan worksheet
        spreadsheet = client.open("arsip_surat")
        worksheet = spreadsheet.sheet1
        
        return worksheet
    except gspread.exceptions.APIError as e:
        if "quota exceeded" in str(e).lower() or "429" in str(e):
            st.error("Quota permintaan ke Google Sheets API telah habis. Silakan coba lagi nanti.")
        else:
            st.error(f"Terjadi kesalahan saat mengakses Google Sheets API: {e}")
        return None

worksheet = access_spreadsheet() # data spreadsheet
def arsip(data):
    '''Menampilkan Arsip Surat dalam SpreadSheet
    dan Mendownload Arsip Surat dari SpreadSheet'''
    data = worksheet.get_all_records()
    st.table(data)
    
def main():
    st.sidebar.title('Aplikasi Pembuat Surat Otomatis')

    surat = st.sidebar.radio('Menu', ('Pemberitahuan', 'Peminjaman', 'Arsip Surat'))
    
    if surat == 'Pemberitahuan':
        srt_pemberitahuan()
    elif surat == 'Peminjaman':
        srt_peminjaman()
    elif surat == 'Arsip Surat':
        arsip(worksheet)
        
    
if __name__ == '__main__':
    main()